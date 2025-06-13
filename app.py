from flask import Flask, request, send_file, jsonify
from flask_cors import CORS
import os, uuid, zipfile
import pandas as pd
from docx import Document
import subprocess


app = Flask(__name__)
CORS(app)

BASE_DIR = os.path.abspath(os.path.dirname(__file__))
UPLOADS = os.path.join(BASE_DIR, "uploads")
RESULTS = os.path.join(BASE_DIR, "results")
os.makedirs(UPLOADS, exist_ok=True)
os.makedirs(RESULTS, exist_ok=True)

def replace_placeholders(doc, data):
    def process_runs(runs):
        i = 0
        while i < len(runs):
            combined = ""
            run_indices = []
            j = i
            while j < len(runs) and len(combined) < 50:  # reasonable max placeholder length
                combined += runs[j].text
                run_indices.append(j)
                for key, value in data.items():
                    placeholder = f"{{{key}}}"
                    if placeholder in combined:
                        replaced_text = combined.replace(placeholder, str(value))
                        # Replace only once
                        runs[i].text = replaced_text
                        # Clear the rest
                        for k in run_indices[1:]:
                            runs[k].text = ""
                        return True  # restart scan
                j += 1
            i += 1
        return False

    # Paragraphs
    for para in doc.paragraphs:
        while process_runs(para.runs):
            pass

    # Tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    while process_runs(para.runs):
                        pass

@app.route("/generate", methods=["POST"])
def generate():
    word_file = request.files["docx"]
    excel_file = request.files["excel"]
    session_id = str(uuid.uuid4())
    session_dir = os.path.join(RESULTS, session_id)
    os.makedirs(session_dir, exist_ok=True)

    word_path = os.path.join(session_dir, "template.docx")
    excel_path = os.path.join(session_dir, "data.xlsx")
    word_file.save(word_path)
    excel_file.save(excel_path)

    df = pd.read_excel(excel_path)
    df.columns = df.columns.str.strip()

    for i, row in df.iterrows():
        data_dict = {col: row[col] for col in df.columns}
        filled_docx = os.path.join(session_dir, f"{i+1}.docx")
        pdf_path = os.path.join(session_dir, f"{i+1}.pdf")

        doc = Document(word_path)
        replace_placeholders(doc, data_dict)
        doc.save(filled_docx)

        subprocess.run([
            "libreoffice",
            "--headless",
            "--convert-to", "pdf",
            "--outdir", session_dir,
            filled_docx
        ])



    zip_path = os.path.join(RESULTS, f"{session_id}.zip")
    with zipfile.ZipFile(zip_path, "w") as zipf:
        for file in os.listdir(session_dir):
            if file.endswith(".pdf"):
                zipf.write(os.path.join(session_dir, file), arcname=file)

    return send_file(zip_path, as_attachment=True, mimetype='application/zip')

if __name__ == "__main__":
    app.run(host="0.0.0.0", port=5000, debug=False)
